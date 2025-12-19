"""DOCX metadata extractor using python-docx."""

from pathlib import Path
from typing import ClassVar

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from metaextract.core.models import DocumentMetadata, ExtractionResult, FileType
from metaextract.extractors.base import MetadataExtractor


class DOCXExtractor(MetadataExtractor):
    """Extract metadata from DOCX files using python-docx."""

    SUPPORTED_TYPES: ClassVar[list[FileType]] = [FileType.DOCX]

    def extract(self) -> ExtractionResult:
        """Extract metadata from DOCX."""
        try:
            metadata = self._create_base_metadata()

            doc = Document(str(self.file_path))
            core_props = doc.core_properties

            # Extract author
            if core_props.author:
                metadata.author = self._clean_string(core_props.author)
                if metadata.author and metadata.author not in metadata.users:
                    metadata.users.append(metadata.author)

            # Extract last modified by
            if core_props.last_modified_by:
                metadata.last_modified_by = self._clean_string(core_props.last_modified_by)
                if metadata.last_modified_by and metadata.last_modified_by not in metadata.users:
                    metadata.users.append(metadata.last_modified_by)

            # Extract dates
            metadata.created = core_props.created
            metadata.modified = core_props.modified

            # Extract application info from app.xml
            self._extract_app_info(metadata)

            # Store raw properties
            metadata.raw = {
                "author": str(core_props.author) if core_props.author else None,
                "last_modified_by": str(core_props.last_modified_by) if core_props.last_modified_by else None,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
                "title": str(core_props.title) if core_props.title else None,
                "subject": str(core_props.subject) if core_props.subject else None,
                "category": str(core_props.category) if core_props.category else None,
                "revision": str(core_props.revision) if core_props.revision else None,
            }

            self._metadata = metadata
            return ExtractionResult(success=True, metadata=metadata)

        except PackageNotFoundError as e:
            return ExtractionResult(success=False, error=f"Invalid DOCX file: {e}")
        except Exception as e:
            return ExtractionResult(success=False, error=str(e))

    def _extract_app_info(self, metadata: DocumentMetadata) -> None:
        """Extract application info from the DOCX package."""
        import zipfile
        from lxml import etree

        try:
            with zipfile.ZipFile(self.file_path, "r") as zf:
                if "docProps/app.xml" in zf.namelist():
                    app_xml = zf.read("docProps/app.xml")
                    root = etree.fromstring(app_xml)

                    # Define namespace
                    ns = {"ep": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"}

                    # Extract application
                    app_elem = root.find("ep:Application", ns)
                    if app_elem is not None and app_elem.text:
                        metadata.application = app_elem.text
                        if app_elem.text not in metadata.software:
                            metadata.software.append(app_elem.text)

                    # Extract app version
                    version_elem = root.find("ep:AppVersion", ns)
                    if version_elem is not None and version_elem.text:
                        metadata.app_version = version_elem.text

                    # Extract template
                    template_elem = root.find("ep:Template", ns)
                    if template_elem is not None and template_elem.text:
                        metadata.template = template_elem.text
                        # Template path might reveal server/path info
                        if "/" in template_elem.text or "\\" in template_elem.text:
                            metadata.paths.append(template_elem.text)

        except Exception:
            pass  # App info is optional

    def extract_text(self) -> str | None:
        """Extract text content for email parsing."""
        try:
            doc = Document(str(self.file_path))
            paragraphs = [para.text for para in doc.paragraphs]
            return "\n".join(paragraphs)
        except Exception:
            return None
